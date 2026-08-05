from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "legal" / "contratos"

NAVY = "122545"
TEXT = "1D1D1F"
MUTED = "5F6368"
LIGHT = "F2F4F7"
NOTICE = "FBF3DF"
WHITE = "FFFFFF"
ACCENT = "1A6B3A"
FONT = "Arial"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("Las columnas deben sumar 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=10.5, bold=False, italic=False, color=TEXT) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 14, 14, 7),
        ("Heading 2", 12, 10, 5),
        ("Heading 3", 11, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_header_footer(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"PQRS SERVICES  |  {short_title}")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    right = p.add_run("    BORRADOR PARA REVISIÓN")
    set_run_font(right, size=8.5, bold=True, color="8A5A00")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("PQRS Services  |  Borrador no autorizado para firma  |  Página ")
    set_run_font(run, size=8, color=MUTED)
    add_field(p, "PAGE")
    run = p.add_run(" de ")
    set_run_font(run, size=8, color=MUTED)
    add_field(p, "NUMPAGES")


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PQRS SERVICES")
    set_run_font(run, size=11, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(subtitle)
    set_run_font(run, size=11.5, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTICE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "BORRADOR DE TRABAJO. No firmar, publicar ni usar con clientes hasta que un abogado colombiano y el contador del prestador revisen su contenido y efectos tributarios."
    )
    set_run_font(run, size=9.5, bold=True, color="7A5A00")


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, italic=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run)


def add_clause(doc: Document, number: int, title: str, paragraphs: list[str], bullets: list[str] | None = None) -> None:
    doc.add_heading(f"{number}. {title}", level=1)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    hdr = table.rows[0]
    hdr.cells[0].text = "Campo"
    hdr.cells[1].text = "Información acordada"
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=WHITE)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_width(cells[0], 2700)
        set_cell_width(cells[1], 6660)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT)
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=NAVY)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5)


def add_signature_table(doc: Document, left_title: str, right_title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [4680, 4680])
    values = [
        (left_title, right_title),
        ("Firma: ______________________________", "Firma: ______________________________"),
        ("Nombre: [___________________________]", "Nombre: [___________________________]"),
        ("Identificación: [____________________]", "Identificación: [____________________]"),
        ("Fecha: [____/____/________]", "Fecha: [____/____/________]"),
    ]
    for row_index, (row, pair) in enumerate(zip(table.rows, values)):
        for idx, value in enumerate(pair):
            row.cells[idx].text = value
            set_cell_width(row.cells[idx], 4680)
            set_cell_margins(row.cells[idx], top=90, bottom=90)
            for run in row.cells[idx].paragraphs[0].runs:
                set_run_font(run, size=9.5, bold=(row_index == 0), color=NAVY if row_index == 0 else TEXT)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT)


def add_review_checklist(doc: Document, items: list[str]) -> None:
    doc.add_page_break()
    doc.add_heading("Anexo de revisión profesional (no contractual)", level=1)
    add_paragraph(
        doc,
        "Esta página existe únicamente para facilitar la revisión del abogado y del contador. Debe retirarse de la versión final de firma.",
        italic=True,
    )
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        run = p.add_run(f"[  ] {item}")
        set_run_font(run, size=10)


def new_document(title: str, subject: str, short_title: str) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc, short_title)
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "PQRS Services"
    props.keywords = "PQRS Services, contrato, borrador, Colombia"
    props.comments = "Borrador sujeto a revisión jurídica y contable"
    return doc


def build_customer_contract() -> Path:
    doc = new_document(
        "Contrato marco de licencia y prestación de servicios",
        "Borrador contractual entre PQRS Services y una propiedad horizontal",
        "CONTRATO MARCO",
    )
    add_title_block(
        doc,
        "Contrato marco de licencia y prestación de servicios",
        "Plataforma PQRS Services para propiedades horizontales",
    )
    add_paragraph(
        doc,
        "Entre [NOMBRE COMPLETO DEL PRESTADOR], identificado con [C.C./NIT] [NÚMERO], con domicilio en [CIUDAD], quien opera bajo la marca PQRS Services y en adelante se denomina EL PRESTADOR; y [NOMBRE LEGAL DE LA COPROPIEDAD], identificada con NIT [NÚMERO], representada legalmente por [NOMBRE], identificado con [DOCUMENTO], en adelante EL CLIENTE, se celebra el presente contrato.",
    )
    add_paragraph(
        doc,
        "El representante de EL CLIENTE declara que cuenta con las facultades legales, estatutarias y presupuestales necesarias para contratar, y entregará el certificado de representación legal y la autorización del órgano competente cuando corresponda.",
    )

    add_clause(doc, 1, "Objeto", [
        "EL PRESTADOR concede a EL CLIENTE una licencia limitada, no exclusiva, no transferible y temporal para usar PQRS Services, y presta los servicios de configuración, alojamiento, soporte y mantenimiento descritos en la Orden de Servicio.",
        "La plataforma permite radicar, gestionar, consultar y reportar peticiones, quejas, reclamos y solicitudes, con acceso diferenciado para administración, consejo y residentes. Los módulos no indicados como incluidos en la Orden de Servicio no forman parte del precio.",
    ])

    add_clause(doc, 2, "Documentos que integran el contrato", [
        "Hacen parte del contrato, en orden de prevalencia: (i) la Orden de Servicio firmada; (ii) este contrato marco; (iii) el Anexo de Tratamiento de Datos; (iv) la autorización de cobro automático, cuando se use; y (v) las políticas públicas vigentes de privacidad y pagos de PQRS Services.",
        "Una condición especial de la Orden de Servicio prevalece únicamente si identifica expresamente la cláusula que modifica.",
    ])

    add_clause(doc, 3, "Inicio, duración y periodos de servicio", [
        "La vigencia inicia en la fecha indicada en la Orden de Servicio. El plazo comercial puede ser de uno o varios años, pero para efectos de precio, renovación, salida y comisión cada año se divide en Periodos Anuales de Servicio de doce (12) meses.",
        "La Orden de Servicio debe indicar si, al finalizar el plazo inicial, existe renovación automática por periodos de doce meses o si la renovación exige una aceptación escrita. Si el campo queda vacío, la renovación no será automática.",
        "La no renovación debe comunicarse con al menos treinta (30) días calendario de anticipación al final del Periodo Anual de Servicio en curso.",
    ])

    add_clause(doc, 4, "Implementación y aceptación", [
        "EL CLIENTE entregará oportunamente su información básica, listado autorizado de usuarios, categorías y demás insumos necesarios. EL PRESTADOR realizará la configuración y acompañamiento pactados en la Orden de Servicio.",
        "La aceptación operativa se documentará mediante acta, correo o registro electrónico. La demora causada por información incompleta de EL CLIENTE ajustará razonablemente el cronograma sin constituir incumplimiento de EL PRESTADOR.",
    ])

    add_clause(doc, 5, "Precio y modalidades de pago", [
        "La tarifa depende del número de unidades, los módulos y las condiciones comerciales de la Orden de Servicio. El plazo contractual no cambia por escoger una forma de pago diferente.",
    ], [
        "Mensual manual: EL CLIENTE paga cada mensualidad por transferencia u otro medio habilitado.",
        "Mensual automático: EL CLIENTE autoriza cobros recurrentes mensuales mediante Wompi o la pasarela indicada.",
        "Anual anticipado: EL CLIENTE paga doce meses por adelantado y recibe un descuento del diez por ciento (10 %) sobre la tarifa mensual de lista multiplicada por doce.",
    ])
    add_paragraph(doc, "Desactivar o revocar el cobro automático no termina el contrato ni extingue las obligaciones pendientes; desde ese momento EL CLIENTE deberá pagar por un medio manual habilitado.")
    add_paragraph(doc, "Los impuestos, retenciones y documentos tributarios se aplicarán conforme a la situación fiscal de las partes y a la normativa vigente. El precio permanecerá fijo durante cada Periodo Anual de Servicio, salvo cambios de unidades, módulos, impuestos o acuerdo escrito.")

    add_clause(doc, 6, "Mora, periodo de gracia y suspensión", [
        "Ante un pago rechazado o vencido, EL PRESTADOR notificará al administrador y concederá el periodo de gracia señalado en la Orden de Servicio, que no será inferior a cinco (5) días calendario.",
        "Vencido el periodo de gracia sin pago, EL PRESTADOR podrá suspender el acceso sin eliminar la información. La suspensión no equivale a terminación y no libera las obligaciones económicas causadas. El acceso se restablecerá cuando el pago sea confirmado.",
    ])

    add_clause(doc, 7, "Terminación anticipada", [
        "Cualquiera de las partes puede terminar por incumplimiento esencial de la otra, siempre que comunique el incumplimiento y conceda diez (10) días hábiles para corregirlo cuando la corrección sea posible.",
        "EL CLIENTE también puede terminar por conveniencia con aviso escrito de treinta (30) días calendario. Si paga mensualmente y recibió en la Orden de Servicio una ventaja sustancial asociada a la permanencia -por ejemplo, implementación bonificada, precio especial o precio protegido-, pagará como compensación el menor valor entre: (i) dos mensualidades vigentes; y (ii) las mensualidades que falten para concluir el Periodo Anual de Servicio en curso. Si la Orden de Servicio no identifica esa ventaja, no se cobra compensación.",
        "En un plan anual anticipado, no se acumula una penalidad adicional: los meses completos usados se recalculan a la tarifa mensual de lista y se devuelve el saldo correspondiente a meses completos no utilizados. Nunca se cobra más de lo recibido por ese periodo.",
        "Si EL CLIENTE termina por un incumplimiento no corregido de EL PRESTADOR, no habrá compensación y se devolverá proporcionalmente el servicio pagado y no prestado. Estas reglas se aplican sin limitar derechos imperativos que correspondan.",
        "Si a la relación le resulta aplicable el Estatuto del Consumidor, cualquier permanencia mínima será expresa, estará asociada a una ventaja sustancial, no superará un (1) año salvo excepción legal y se ofrecerá una alternativa sin permanencia cuando la ley lo exija.",
    ])

    add_clause(doc, 8, "Obligaciones de EL PRESTADOR", ["EL PRESTADOR se obliga a:"], [
        "Operar la plataforma con medidas razonables de seguridad, continuidad y respaldo.",
        "Respetar los roles y el aislamiento de información entre conjuntos.",
        "Prestar el soporte y acompañamiento descritos en la Orden de Servicio.",
        "Informar mantenimientos programados que afecten materialmente el uso.",
        "Permitir la consulta y exportación de información conforme al alcance contratado.",
        "Tratar los datos personales únicamente para prestar el servicio y según las instrucciones documentadas de EL CLIENTE.",
    ])

    add_clause(doc, 9, "Obligaciones de EL CLIENTE", ["EL CLIENTE se obliga a:"], [
        "Pagar completa y oportunamente el precio acordado.",
        "Designar usuarios autorizados, asignar roles y desactivar accesos que ya no correspondan.",
        "Mantener actualizados sus datos de contacto, unidades y facturación.",
        "Obtener las autorizaciones de tratamiento de datos de residentes y demás titulares.",
        "No cargar contenido ilegal ni datos sensibles o de menores sin base jurídica y autorización expresa.",
        "No intentar acceder a otros conjuntos, alterar controles o realizar extracción masiva no autorizada.",
        "Responder por la veracidad y licitud de la información que sus usuarios registren.",
    ])

    add_clause(doc, 10, "Disponibilidad, soporte y cambios", [
        "Salvo que la Orden de Servicio incluya un acuerdo de nivel de servicio, no se garantiza un porcentaje específico de disponibilidad. EL PRESTADOR realizará esfuerzos razonables para mantener el servicio y atender incidentes.",
        "Los cambios materiales de funcionalidad o tarifa se comunicarán con al menos treinta (30) días calendario de anticipación y se aplicarán desde el siguiente Periodo Anual de Servicio, salvo que sean necesarios por seguridad, cumplimiento legal o corrección urgente.",
    ])

    add_clause(doc, 11, "Datos personales y confidencialidad", [
        "Para los datos de residentes y demás usuarios, EL CLIENTE actúa como Responsable del Tratamiento y EL PRESTADOR como Encargado, conforme a la Ley 1581 de 2012 y sus normas reglamentarias. El Anexo de Tratamiento de Datos contiene las instrucciones, finalidades, medidas y subencargados autorizados.",
        "Cada parte guardará reserva sobre la información no pública de la otra durante el contrato y por dos (2) años después de su terminación; la confidencialidad de datos personales y secretos empresariales subsistirá durante el término exigido por la ley o mientras mantengan ese carácter.",
    ])

    add_clause(doc, 12, "Propiedad intelectual y titularidad de la información", [
        "El software, marca, código, diseño y documentación de PQRS Services pertenecen a EL PRESTADOR o a sus licenciantes. EL CLIENTE conserva la titularidad de la información que carga y de los documentos que genera con sus datos.",
        "El contrato no transfiere código fuente ni concede derechos de reventa, sublicencia, ingeniería inversa o explotación distinta al uso interno del conjunto.",
    ])

    add_clause(doc, 13, "Responsabilidad", [
        "Cada parte responde por sus incumplimientos directos y comprobados. En la medida permitida por la ley, la responsabilidad acumulada de EL PRESTADOR se limita a las sumas efectivamente pagadas por EL CLIENTE durante los tres (3) meses anteriores al hecho que originó la reclamación.",
        "EL PRESTADOR no responde por decisiones administrativas tomadas por EL CLIENTE, contenido ingresado por sus usuarios, fallas de conectividad del cliente, uso indebido de credenciales ni eventos de fuerza mayor. El límite no aplica cuando la ley prohíba limitar la responsabilidad, incluidos el dolo y la culpa grave.",
    ])

    add_clause(doc, 14, "Finalización y entrega de datos", [
        "Terminada la relación, EL CLIENTE contará con treinta (30) días calendario para solicitar la exportación disponible de su información. Cumplido ese plazo, EL PRESTADOR podrá eliminarla o anonimizarla, salvo registros cuya conservación sea legalmente obligatoria.",
    ])

    add_clause(doc, 15, "Notificaciones, firma y evidencia", [
        "Las comunicaciones contractuales se enviarán a los correos indicados en la Orden de Servicio. Cada parte debe informar sus cambios de contacto.",
        "El contrato puede firmarse física o electrónicamente. Las partes aceptan como evidencia los documentos firmados, mensajes de datos, registros de aceptación, sellos de tiempo, direcciones IP y trazas de auditoría que permitan identificar al firmante y preservar la integridad del contenido.",
    ])

    add_clause(doc, 16, "Ley aplicable y controversias", [
        "El contrato se rige por las leyes de Colombia. Las partes intentarán una negociación directa durante treinta (30) días calendario desde el reclamo escrito. Si no hay acuerdo, acudirán a los jueces competentes de [CIUDAD], sin perjuicio de mecanismos de conciliación o derechos imperativos aplicables.",
    ])

    doc.add_heading("Orden de Servicio No. [____]", level=1)
    add_key_value_table(doc, [
        ("Cliente", "[Nombre legal de la copropiedad y NIT]"),
        ("Representante", "[Nombre, documento, cargo y soporte de facultades]"),
        ("Unidades", "[Cantidad de unidades privadas]"),
        ("Módulos incluidos", "[PQRS / Reservas / Pagos de residentes / Otros]"),
        ("Implementación", "[Estándar / asistida / bonificada] y valor $[____]"),
        ("Ventaja por permanencia", "[Describir beneficio sustancial o indicar N/A]"),
        ("Fecha de inicio", "[AAAA-MM-DD]"),
        ("Duración inicial", "[1 / 2 / 3 / __] años"),
        ("Periodo anual", "Doce (12) meses"),
        ("Tarifa mensual de lista", "$[____] COP"),
        ("Forma de pago", "[  ] Mensual manual   [  ] Mensual automática   [  ] Anual anticipada con 10 %"),
        ("Precio efectivo", "$[____] COP por [mes/año]"),
        ("Renovación", "[  ] Automática por 12 meses   [  ] Requiere aceptación escrita"),
        ("Aviso de no renovación", "Treinta (30) días calendario"),
        ("Gracia", "[____] días calendario, mínimo 5"),
        ("Soporte", "[Canal, horario y tiempos objetivo]"),
        ("Contacto administrativo", "[Nombre, correo y teléfono]"),
        ("Contacto de facturación", "[Nombre y correo]"),
        ("Condiciones especiales", "[____]"),
    ])

    doc.add_heading("Anexo A. Tratamiento de datos personales", level=1)
    add_paragraph(doc, "Objeto del encargo. EL PRESTADOR tratará por cuenta de EL CLIENTE los datos necesarios para autenticar usuarios, gestionar PQRS, evidencias, notificaciones, reportes, auditoría y soporte.")
    add_key_value_table(doc, [
        ("Titulares", "Residentes, propietarios, tenedores, trabajadores, contratistas y usuarios autorizados del conjunto."),
        ("Datos", "Identificación y contacto, unidad o vínculo con el conjunto, credenciales cifradas, contenido de PQRS, evidencias, actividad y auditoría."),
        ("Finalidades", "Prestar, asegurar, soportar y documentar el servicio contratado; cumplir obligaciones legales."),
        ("Duración", "Durante la relación contractual y los periodos legales de conservación."),
        ("Subencargados previstos", "Supabase, Vercel, Resend, Wompi y, cuando corresponda, Mercado Pago, o sus reemplazos informados previamente."),
        ("Transferencias", "Las necesarias para operar proveedores tecnológicos, sujetas a garantías legales y contractuales aplicables."),
    ])
    add_bullets(doc, [
        "EL PRESTADOR seguirá instrucciones documentadas, aplicará controles de acceso, cifrado en tránsito, almacenamiento privado y auditoría.",
        "EL PRESTADOR notificará a EL CLIENTE incidentes confirmados sin dilación indebida y, como objetivo contractual, dentro de las 72 horas siguientes a su confirmación.",
        "EL PRESTADOR apoyará consultas, correcciones, supresiones y exportaciones en la medida técnica correspondiente a su rol de Encargado.",
        "EL CLIENTE conservará las autorizaciones, informará finalidades y no suministrará datos sin base jurídica.",
        "Al terminar, los datos se exportarán, eliminarán o anonimizarán conforme a la cláusula 14.",
    ])

    doc.add_heading("Anexo B. Autorización de cobro automático", level=1)
    add_paragraph(doc, "Este anexo solo se firma si EL CLIENTE elige cobro automático. Autoriza a EL PRESTADOR y a Wompi a iniciar cargos recurrentes por el valor y frecuencia indicados en la Orden de Servicio, usando el medio de pago tokenizado por la pasarela.")
    add_bullets(doc, [
        "La plataforma mostrará el valor, frecuencia y fecha estimada del siguiente cobro.",
        "EL CLIENTE puede revocar esta autorización desde Licencias y pagos o mediante comunicación escrita.",
        "La revocatoria detiene cargos futuros, pero no termina el contrato ni elimina obligaciones ya causadas.",
        "PQRS Services no almacena el número completo, código de seguridad ni credenciales bancarias de la tarjeta.",
    ])

    doc.add_heading("Firmas", level=1)
    add_paragraph(doc, "Las partes declaran que leyeron y comprendieron el contrato, la Orden de Servicio y sus anexos.")
    add_signature_table(doc, "EL PRESTADOR", "EL CLIENTE")

    add_review_checklist(doc, [
        "Verificar identidad, RUT, obligaciones de facturación, IVA y retenciones de EL PRESTADOR.",
        "Validar facultades del representante del conjunto conforme a Ley 675 de 2001 y al reglamento de propiedad horizontal.",
        "Confirmar si el Estatuto del Consumidor aplica y ajustar permanencia, retracto, reversión y compensación de salida.",
        "Revisar que la compensación máxima de dos mensualidades sea válida, proporcionada y soportada por una ventaja sustancial identificada.",
        "Definir si contratos plurianuales serán marcos con salidas anuales o compromisos no cancelables negociados individualmente.",
        "Revisar límite de responsabilidad de tres meses, cobertura de seguros y régimen de daños.",
        "Confirmar conservación, exportación y eliminación de datos; revisar subencargados y transferencias internacionales.",
        "Alinear la Orden de Servicio con precios, módulos, Wompi, renovaciones y registros reales de la plataforma.",
        "Eliminar esta página y la marca de borrador antes de la firma final.",
    ])

    output = OUTPUT_DIR / "Contrato_marco_servicios_PQRS_Services_BORRADOR.docx"
    doc.save(output)
    return output


def build_referral_agreement() -> Path:
    doc = new_document(
        "Acuerdo independiente de referidos y gestión comercial",
        "Borrador de acuerdo de comisiones para referidos de PQRS Services",
        "ACUERDO DE REFERIDOS",
    )
    add_title_block(
        doc,
        "Acuerdo independiente de referidos y gestión comercial",
        "Reglas de atribución, comisión, renovación y pago gradual",
    )
    add_paragraph(
        doc,
        "Entre [NOMBRE COMPLETO DEL PRESTADOR], identificado con [C.C./NIT] [NÚMERO], titular de la operación comercial de PQRS Services, en adelante EL PRESTADOR; y [NOMBRE COMPLETO DE LA GESTORA], identificada con C.C. [NÚMERO], en adelante LA GESTORA, se celebra este acuerdo independiente.",
    )

    add_clause(doc, 1, "Objeto", [
        "LA GESTORA podrá identificar y presentar propiedades horizontales potencialmente interesadas en PQRS Services, facilitar contactos y apoyar conversaciones comerciales. EL PRESTADOR conserva en todo momento la decisión sobre precios, demostraciones, propuestas, contratos y aceptación de clientes.",
    ])

    add_clause(doc, 2, "Naturaleza independiente", [
        "Este acuerdo no crea contrato de trabajo, sociedad, mandato, representación, franquicia ni exclusividad. LA GESTORA organiza autónomamente su actividad, no está sujeta a horario, cuota, reglamento, supervisión continua ni órdenes sobre modo, tiempo o cantidad de trabajo.",
        "LA GESTORA no recibe salario, prestaciones ni reembolso de gastos salvo autorización escrita. La única contraprestación es la comisión causada conforme a este acuerdo.",
        "Las partes no asignan territorio ni un encargo estable de explotación comercial. Si la operación real cambia y adquiere características de agencia comercial o relación laboral, deberán renegociar el instrumento antes de continuar.",
    ])

    add_clause(doc, 3, "Registro y atribución de referidos", [
        "Antes de presentar una propuesta, LA GESTORA registrará por correo o en el formato anexo el nombre del conjunto, contacto autorizado, fuente del contacto y fecha. EL PRESTADOR confirmará por escrito si el referido es aceptado.",
    ], [
        "No se atribuye comisión por un cliente ya activo, un prospecto con negociación documentada en los seis meses anteriores o un contacto obtenido directamente por EL PRESTADOR antes del registro.",
        "La atribución inicial dura ciento veinte (120) días, prorrogables por escrito cuando exista una negociación activa.",
        "Si dos personas reclaman el mismo referido, prevalece el primer registro aceptado; cualquier distribución exige acuerdo escrito.",
        "La amistad, cercanía o relación profesional con un contacto no invalida el referido, siempre que no exista conflicto de interés ni pago oculto al decisor.",
    ])

    add_clause(doc, 4, "Límites de actuación", ["LA GESTORA no está autorizada para:"], [
        "Firmar contratos, aceptar obligaciones o representar a PQRS Services.",
        "Modificar precios, descuentos, permanencia, soporte o funcionalidades.",
        "Recibir pagos de clientes en nombre de EL PRESTADOR.",
        "Prometer desarrollos, integraciones o plazos no aprobados por escrito.",
        "Compartir accesos, información confidencial o datos personales de residentes.",
    ])

    add_clause(doc, 5, "Fórmula de la comisión", [
        "Por cada año completo contratado y efectivamente activado de un referido aceptado, LA GESTORA tendrá derecho a una comisión equivalente a una (1) mensualidad neta recurrente del año correspondiente.",
        "La mensualidad neta recurrente es el valor de suscripción efectivamente acordado y recibido, después de descuentos, dividido por los meses cubiertos. En un pago anual se calcula como valor anual neto dividido por doce (12). Se excluyen IVA, retenciones, implementación, módulos de pago único, gastos reembolsables, costos de pasarela, intereses, devoluciones y cortesías.",
        "Ejemplos: un contrato de un año genera una mensualidad; uno de dos años genera dos mensualidades; uno de cinco años genera cinco mensualidades, pagaderas gradualmente según la cláusula 6.",
    ])

    add_clause(doc, 6, "Causación y pago gradual", [
        "La comisión total se divide en una cuota por cada Periodo Anual de Servicio. Ninguna cuota futura se causa únicamente porque el cliente haya firmado un contrato plurianual o pagado anticipadamente varios años.",
    ], [
        "Primera cuota con pago mensual: se causa cuando el cliente ha pagado dos mensualidades aprobadas y no reversadas.",
        "Primera cuota con pago anual: se causa cuando el pago anual está aprobado, han transcurrido treinta (30) días calendario sin devolución total y el servicio está activo.",
        "Cuotas de años posteriores: se causan en cada aniversario, siempre que el Periodo Anual de Servicio correspondiente haya empezado, el contrato continúe vigente y el pago de ese año esté aprobado.",
        "Si el cliente paga varios años por anticipado, las cuotas continúan pagándose en cada aniversario, no todas al inicio.",
        "Cada cuota causada se pagará dentro de los diez (10) días hábiles siguientes a la entrega de la cuenta de cobro o documento tributario exigible y los datos bancarios actualizados.",
    ])

    add_clause(doc, 7, "Renovaciones", [
        "Una renovación genera la misma comisión de una mensualidad por cada año adicional únicamente cuando: (i) se firma o acepta la renovación; (ii) el cliente paga el periodo correspondiente; y (iii) LA GESTORA tuvo una intervención comercial material y documentada, aceptada por EL PRESTADOR antes de la renovación.",
        "Una renovación automática, por sí sola, no genera comisión. Tampoco la genera una renovación gestionada directamente por EL PRESTADOR sin intervención demostrable de LA GESTORA.",
    ])

    add_clause(doc, 8, "Cancelaciones, mora y reversos", [
        "Si el cliente termina, no renueva, entra en mora definitiva o recibe devolución antes de iniciar un Periodo Anual de Servicio, las cuotas futuras no causadas se extinguen.",
        "Las comisiones correctamente pagadas no se devuelven, salvo fraude, duplicidad, falsedad atribuible a LA GESTORA o reversión del pago que sirvió como base. En esos casos EL PRESTADOR podrá compensar el valor contra comisiones futuras, previa entrega de soporte.",
        "La terminación anticipada del contrato del cliente no acelera las cuotas futuras.",
    ])

    add_clause(doc, 9, "Conflictos de interés e integridad", [
        "LA GESTORA informará cualquier relación económica, laboral, contractual, familiar o de representación que pueda influir en la decisión del conjunto referido.",
        "Si LA GESTORA llega a ser administradora, consejera, empleada, contratista decisora o representante del conjunto prospecto, la comisión solo podrá mantenerse con revelación escrita y autorización del órgano competente del conjunto. Si no existe esa autorización, el referido no genera comisión.",
        "LA GESTORA no ofrecerá regalos, pagos, participaciones de comisión ni beneficios ocultos a administradores, consejeros, revisores fiscales, empleados o decisores del cliente. Cualquier práctica de este tipo permite terminar el acuerdo y cancelar comisiones no pagadas.",
    ])

    add_clause(doc, 10, "Confidencialidad y datos personales", [
        "LA GESTORA usará la información comercial únicamente para presentar el servicio y guardará confidencialidad durante el acuerdo y por dos (2) años después. No solicitará bases de residentes, credenciales ni información de PQRS.",
        "Los datos de contacto de prospectos solo podrán compartirse cuando exista autorización o una base jurídica aplicable, y se limitarán a lo necesario para la conversación comercial.",
    ])

    add_clause(doc, 11, "Impuestos y soportes", [
        "Cada parte es responsable de sus obligaciones tributarias. EL PRESTADOR practicará las retenciones legalmente aplicables y entregará los certificados correspondientes cuando tenga esa obligación.",
        "LA GESTORA entregará RUT actualizado, certificación bancaria y cuenta de cobro, factura o documento equivalente que corresponda. El plazo de pago empieza cuando la documentación esté completa.",
    ])

    add_clause(doc, 12, "Información y verificación", [
        "EL PRESTADOR mantendrá un registro por referido con contrato, duración, modalidad de pago, estado de cada cuota, valor base, fecha de causación y referencia de pago. LA GESTORA puede solicitar un estado razonable de sus referidos sin acceder a información confidencial del cliente.",
        "Un desacuerdo sobre el cálculo debe informarse dentro de los treinta (30) días calendario siguientes al reporte o pago; las partes revisarán los soportes pertinentes.",
    ])

    add_clause(doc, 13, "Duración y terminación del acuerdo", [
        "El acuerdo tendrá una duración inicial de doce (12) meses y se renovará por periodos iguales salvo aviso escrito con treinta (30) días calendario de anticipación. Cualquiera de las partes puede terminarlo por conveniencia con el mismo aviso.",
        "La terminación no afecta cuotas ya causadas ni referidos aceptados con contrato vigente; sus cuotas futuras seguirán las cláusulas 6 a 8, salvo terminación por fraude, soborno, violación de confidencialidad o conflicto de interés no revelado.",
    ])

    add_clause(doc, 14, "Ley aplicable y controversias", [
        "El acuerdo se rige por las leyes de Colombia. Las partes intentarán resolver directamente cualquier diferencia durante treinta (30) días calendario y, si no lo logran, acudirán a los jueces competentes de [CIUDAD], sin perjuicio de una conciliación voluntaria.",
    ])

    doc.add_heading("Anexo 1. Registro de referido", level=1)
    add_key_value_table(doc, [
        ("Código", "[REF-____]"),
        ("Fecha de registro", "[AAAA-MM-DD]"),
        ("Conjunto", "[Nombre legal y nombre comercial]"),
        ("NIT", "[____]"),
        ("Contacto", "[Nombre, cargo, correo y teléfono]"),
        ("Origen y relación", "[Cómo se obtuvo el contacto y relación existente]"),
        ("Conflicto revelado", "[Ninguno / descripción y autorización adjunta]"),
        ("Estado previo", "[Nuevo / ya conocido / negociación existente]"),
        ("Aceptación del referido", "[Aceptado / rechazado] por [____] el [____]"),
        ("Vigencia de atribución", "Hasta [AAAA-MM-DD]"),
        ("Contrato conseguido", "[Número, fecha y duración]"),
        ("Modalidad", "[Mensual manual / mensual automática / anual]"),
        ("Valor mensual neto", "$[____] COP"),
        ("Años comisionables", "[____]"),
        ("Calendario de cuotas", "[Fechas estimadas y condiciones]"),
        ("Notas", "[____]"),
    ])

    doc.add_heading("Anexo 2. Registro de cuota de comisión", level=1)
    add_key_value_table(doc, [
        ("Referido", "[REF-____]"),
        ("Periodo anual", "Año [____] de [____]"),
        ("Fecha de inicio del periodo", "[AAAA-MM-DD]"),
        ("Pago del cliente verificado", "[Referencia y fecha]"),
        ("Mensualidad neta base", "$[____] COP"),
        ("Retenciones", "$[____] COP"),
        ("Valor neto pagado", "$[____] COP"),
        ("Fecha de causación", "[AAAA-MM-DD]"),
        ("Fecha y referencia de pago", "[AAAA-MM-DD / referencia]"),
        ("Estado", "[Pendiente / causada / pagada / cancelada]"),
    ])

    doc.add_heading("Firmas", level=1)
    add_paragraph(doc, "Las partes declaran que comprendieron la fórmula, el pago gradual y la independencia de la relación.")
    add_signature_table(doc, "EL PRESTADOR", "LA GESTORA")

    add_review_checklist(doc, [
        "Confirmar que la operación real no configure subordinación laboral según el artículo 23 del Código Sustantivo del Trabajo.",
        "Revisar riesgo de agencia comercial por estabilidad, territorio, promoción continua o representación.",
        "Validar fórmula de una mensualidad neta por año, causación gradual y tratamiento de pagos anticipados plurianuales.",
        "Definir retenciones, IVA si corresponde, cuenta de cobro/factura y obligaciones de información exógena.",
        "Revisar supervivencia de cuotas futuras tras terminar el acuerdo y causales que permiten cancelarlas.",
        "Confirmar protocolo de conflicto de interés si LA GESTORA adquiere un cargo o poder de decisión en un conjunto referido.",
        "Alinear los campos con el modelo de datos y la trazabilidad de comisiones de la plataforma.",
        "Eliminar esta página y la marca de borrador antes de la firma final.",
    ])

    output = OUTPUT_DIR / "Acuerdo_referidos_gestion_comercial_BORRADOR.docx"
    doc.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_customer_contract(), build_referral_agreement()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
